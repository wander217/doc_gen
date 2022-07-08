from docx import Document
from python_docx_replace.docx_replace import docx_replace
from python_docx_replace.docx_replace import MaxRetriesReached

paragraph_datas = {
    "city_name": "THÀNH PHỐ HÀ NỘI",
    "contract_code": "S1232322",
    "first_date": "Đăng ký lần đầu, Thứ 6 ngày 8 tháng 7 năm 2022",
    "change_date": "Đăng ký lần thứ 2, Thứ 6 ngày 8 tháng 7 năm 2022",
    "vietnamese_name": "Công ty cổ phần ABC",
    "foreign_name": "Công ty ABC",
    "acronym_name": "ABC company",
    "phone": "0328296433",
    "fax": "328296433",
    "email": "abc@gmail.com.vn",
    "website": "abc.com.vn",
    "capital_amount": "1000 đồng",
    "capital_value": "100 đồng",
    "total_capital": "100 đồng",
    "own_fullname": "nguyễn văn A",
    "own_sex": "nam",
    "own_position": "abcde",
    "own_birthday": "21/07/1999",
    "own_ethnicity": "Kinh",
    "own_nation": "Việt Nam",
    "own_type": "Chứng minh thư nhân dân",
    "own_idcard_code": "0123456788",
    "own_idcard_date": "12/03/2010",
    "own_idcard_place": "Công an Hà Nội",
    "own_idcard_rp": "Số 1, đường A, phố B, khu C, Thành phố D",
    "own_idcard_lv": "Số 1, đường A, phố B, khu C, Thành phố D",
}

table_datas = {
    "bt": [
        {"bt_id": "0", "bt_name": "Abc", "bt_code": "123"},
        {"bt_id": "1", "bt_name": "Def", "bt_code": "345"},
    ]
}


def get_table_from_keyword(keyword, document):
    for table in document.tables:
        for row in table.rows:
            for col in row.cells:
                for paragraph in col.paragraphs:
                    if keyword in paragraph.text:
                        return table


def _simple_replace(p, key, value):
    """Try to replace a key in the paragraph runs, simpler alternative"""
    for run in p.runs:
        if key in run.text:
            run.text = run.text.replace(f"${{{key}}}", value)


class RunTextChanger:
    def __init__(self, p, key, value):
        self.p = p
        self.key = key
        self.value = value
        self.run_text = ""
        self.runs_indexes = []
        self.run_char_indexes = []
        self.runs_to_change = {}

    def _initialize(self):
        run_index = 0
        for run in self.p.runs:
            self.run_text += run.text
            self.runs_indexes += [run_index for _ in run.text]
            self.run_char_indexes += [char_index for char_index, char in enumerate(run.text)]
            run_index += 1

    def replace(self):
        self._initialize()
        parsed_key_length = len(self.key)
        index_to_replace = self.run_text.find(self.key)

        for i in range(parsed_key_length):
            index = index_to_replace + i
            run_index = self.runs_indexes[index]
            run = self.p.runs[run_index]
            run_char_index = self.run_char_indexes[index]

            if not self.runs_to_change.get(run_index):
                self.runs_to_change[run_index] = [char for char_index, char in enumerate(run.text)]

            run_to_change = self.runs_to_change.get(run_index)
            if index == index_to_replace:
                run_to_change[run_char_index] = self.value
            else:
                run_to_change[run_char_index] = ""

        # make the real replace
        for index, text in self.runs_to_change.items():
            run = self.p.runs[index]
            run.text = "".join(text)


def _complex_replace(p, key, value):
    """Complex alternative, which check all broken items inside the runs"""
    max_retries_replace_a_key = 100  # to avoid infinite loop, this value is set

    changer = RunTextChanger(p, key, value)
    changer.replace()

    current = 1
    while key in p.text:  # if the key appears more than once in the paragraph, it will replaced all
        current += 1
        if current > max_retries_replace_a_key:
            raise MaxRetriesReached(max_retries_replace_a_key, key)
        changer = RunTextChanger(p, key, value)
        changer.replace()


def processing_table_data(key, value, document):
    table = get_table_from_keyword(key, document)
    first_row = table.rows[1]
    for i, row_data in enumerate(value):
        business_path: str = r'D:\python_project\doc_gen\template\business_type.docx'
        business_doc: Document = Document(docx=business_path)
        temp_row = business_doc.tables[0].rows[1]
        for col in temp_row.cells:
            for paragraph in col.paragraphs:
                for k, v in row_data.items():
                    k1 = f"${{{k}}}"
                    if k1 in paragraph.text:
                        _simple_replace(paragraph, k1, v)
                        if k1 in paragraph.text:
                            _complex_replace(paragraph, k1, v)
        last_row = table.rows[-1]
        last_row._tr.addnext(temp_row._tr)
    table._tbl.remove(first_row._tr)


if __name__ == "__main__":
    template_path: str = r'D:\python_project\doc_gen\template\cong_ty_co_phan.docx'
    document: Document = Document(docx=template_path)
    docx_replace(document, **paragraph_datas)
    for key, value in table_datas.items():
        processing_table_data(key, value, document)
    save_path = r'D:\python_project\doc_gen\result\save_doc.docx'
    document.save(save_path)
